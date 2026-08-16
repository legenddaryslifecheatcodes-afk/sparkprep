"""Reads an uploaded .docx manuscript and converts it into the plain-text
chapter format manuscript_composer.py already expects (# Chapter Title
markers). This is deliberately NOT a general-purpose DOCX-to-PDF
converter -- it doesn't try to preserve the original Word file's visual
formatting (fonts, colors, custom styles). Instead it extracts the
*content structure* (which paragraphs are chapter headings vs. body
text) and lets the existing, already-correct interior typesetting engine
(margins, trim size, PDF/X-1a output) lay it out properly for print.

Why this approach instead of a literal Word-to-PDF conversion: most
self-published authors' Word formatting isn't print-ready anyway (wrong
margins, screen-sized fonts, no proper page breaks) -- converting it
as-is would just produce a non-compliant PDF instead of a compliant one.
Re-typesetting the actual text content through manuscript_composer.py
is what actually gets them a submittable interior.

Embedded images aren't placed into the composed interior (there's no
per-page layout editor yet to position them), but they aren't silently
discarded either -- extract_embedded_images() pulls every image out of
the .docx as its own file so the project keeps them, and a person can
add them back through the cover/asset upload flow once a real interior
image-placement editor exists.
"""
import io
import zipfile
from typing import Optional
import docx


HEADING_STYLE_PREFIXES = ("Heading 1", "Title", "Chapter")


def _is_chapter_heading(paragraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    if style_name.startswith(HEADING_STYLE_PREFIXES):
        return True
    # Some authors bold a short standalone line as a de facto heading
    # instead of using Word's built-in Heading style. Treat a short,
    # fully-bold paragraph as a heading too, since manuscript_composer.py
    # only understands '# Title' markers, not bold-run detection.
    text = paragraph.text.strip()
    if not text or len(text) > 80:
        return False
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def extract_manuscript_text(file_like) -> dict:
    """Reads a .docx file (path or file-like object) and returns:
        {"source_text": "<# Chapter markers>...", "chapter_count": N,
         "paragraph_count": N, "word_count": N, "warnings": [...]}
    source_text is formatted exactly as manuscript_composer.parse_source_text
    expects, so it can be passed straight into compose_manuscript_pdf.
    """
    document = docx.Document(file_like)
    lines = []
    chapter_count = 0
    paragraph_count = 0
    word_count = 0
    warnings = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            lines.append("")
            continue
        if _is_chapter_heading(paragraph):
            lines.append(f"# {text}")
            chapter_count += 1
        else:
            lines.append(text)
            paragraph_count += 1
            word_count += len(text.split())

    if chapter_count == 0 and paragraph_count > 0:
        warnings.append(
            "No chapter headings were detected (no text styled with Word's "
            "'Heading 1' or similar). The whole manuscript will be treated "
            "as a single unbroken chapter. If you have chapter breaks, "
            "style each chapter title with Word's Heading 1 style and "
            "re-upload for proper chapter-start pages."
        )
    if document.inline_shapes:
        warnings.append(
            f"This document contains {len(document.inline_shapes)} embedded "
            "image(s), which were extracted separately (not placed into the "
            "composed interior yet -- there's no page-image-placement editor "
            "built for that). You'll need to add them back manually once "
            "that exists."
        )
    if any(t.rows for t in document.tables):
        warnings.append(
            "This document contains one or more tables, which are not "
            "carried over -- only paragraph text is extracted."
        )

    return {
        "source_text": "\n".join(lines),
        "chapter_count": chapter_count,
        "paragraph_count": paragraph_count,
        "word_count": word_count,
        "warnings": warnings,
    }


def extract_embedded_images(file_bytes: bytes) -> list[dict]:
    """Pulls every embedded image out of a .docx and returns them as raw
    bytes with a filename and content type, so they can be saved to the
    project instead of silently disappearing during text extraction.

    A .docx is itself a zip archive; embedded images live under
    word/media/ regardless of whether they're anchored inline, floating,
    or in a header/footer, so reading the zip directly catches every
    image python-docx's higher-level inline_shapes API can miss (e.g.
    floating/wrapped images).
    """
    images = []
    content_types = {
        "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
        "emf": "image/x-emf", "wmf": "image/x-wmf",
    }
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            images.append({
                "filename": name.rsplit("/", 1)[-1],
                "content_type": content_types.get(ext, "application/octet-stream"),
                "data": z.read(name),
            })
    return images
